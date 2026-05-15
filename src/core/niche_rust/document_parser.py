"""
Zenic-Agents — Document Parser (Phase 6.B)

Python-side document parsing for formats that require Python
libraries (PDF, DOCX, HTML). The Rust ingestion engine handles
TXT, CSV, JSON, and Markdown natively; this module fills the
gap for binary formats.

Supported Formats:
    - PDF: PyPDF2 / pdfminer.six (fallback)
    - DOCX: python-docx
    - HTML: built-in html.parser
    - TXT/CSV/JSON: delegated to Rust (ingest_extract_text_simple)

Fallback:
    If a Python library is not installed, the parser returns an
    empty string with an error message. It never raises exceptions.
"""

from __future__ import annotations

import logging
import os
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────────────────────
#  Constants
# ──────────────────────────────────────────────────────────────

#: Maximum file size for Python-side parsing (50 MB).
MAX_FILE_SIZE: int = 50 * 1024 * 1024

#: Supported format strings for Python parsing.
PYTHON_FORMATS: Tuple[str, ...] = ("pdf", "docx", "html")

#: Format string for unknown/unparseable documents.
FORMAT_UNKNOWN: str = "unknown"


# ──────────────────────────────────────────────────────────────
#  PDF Parser
# ──────────────────────────────────────────────────────────────

def _parse_pdf(file_path: str) -> Tuple[str, List[str]]:
    """Extract text from a PDF file using Python libraries.

    Tries PyPDF2 first, then falls back to pdfminer.six.
    Returns (text, errors) tuple — never raises.

    Args:
        file_path: Absolute path to the PDF file.

    Returns:
        Tuple of (extracted_text, error_messages).
    """
    errors: List[str] = []
    text_parts: List[str] = []

    # Try PyPDF2
    try:
        from PyPDF2 import PdfReader  # type: ignore[import-untyped]

        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        if text_parts:
            return ("\n\n".join(text_parts), errors)
    except ImportError:
        errors.append("PyPDF2 not installed, trying pdfminer.six")
    except Exception as e:
        errors.append(f"PyPDF2 error: {e}")

    # Try pdfminer.six
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract  # type: ignore[import-untyped]

        text = pdfminer_extract(file_path)
        if text and text.strip():
            return (text, errors)
    except ImportError:
        errors.append("pdfminer.six not installed")
    except Exception as e:
        errors.append(f"pdfminer.six error: {e}")

    # No parser available
    if not text_parts:
        errors.append(
            "No PDF parser available. Install PyPDF2 or pdfminer.six: "
            "pip install PyPDF2 pdfminer.six"
        )

    return ("", errors)


def _parse_pdf_bytes(data: bytes) -> Tuple[str, List[str]]:
    """Extract text from PDF bytes.

    Writes bytes to a temporary file and delegates to _parse_pdf.

    Args:
        data: PDF file bytes.

    Returns:
        Tuple of (extracted_text, error_messages).
    """
    import tempfile

    errors: List[str] = []
    tmp_path: Optional[str] = None

    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        with open(tmp_path, "wb") as f:
            f.write(data)
        return _parse_pdf(tmp_path)
    except Exception as e:
        errors.append(f"PDF bytes parse error: {e}")
        return ("", errors)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


# ──────────────────────────────────────────────────────────────
#  DOCX Parser
# ──────────────────────────────────────────────────────────────

def _parse_docx(file_path: str) -> Tuple[str, List[str]]:
    """Extract text from a DOCX file using python-docx.

    Args:
        file_path: Absolute path to the DOCX file.

    Returns:
        Tuple of (extracted_text, error_messages).
    """
    errors: List[str] = []
    text_parts: List[str] = []

    try:
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(file_path)

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_parts: List[str] = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_parts.append(cell.text.strip())
                if row_parts:
                    text_parts.append(" | ".join(row_parts))

        return ("\n".join(text_parts), errors)

    except ImportError:
        errors.append(
            "python-docx not installed. Install with: pip install python-docx"
        )
        return ("", errors)
    except Exception as e:
        errors.append(f"python-docx error: {e}")
        return ("", errors)


def _parse_docx_bytes(data: bytes) -> Tuple[str, List[str]]:
    """Extract text from DOCX bytes.

    Args:
        data: DOCX file bytes.

    Returns:
        Tuple of (extracted_text, error_messages).
    """
    import tempfile

    errors: List[str] = []
    tmp_path: Optional[str] = None

    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        with open(tmp_path, "wb") as f:
            f.write(data)
        return _parse_docx(tmp_path)
    except Exception as e:
        errors.append(f"DOCX bytes parse error: {e}")
        return ("", errors)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


# ──────────────────────────────────────────────────────────────
#  HTML Parser
# ──────────────────────────────────────────────────────────────

def _parse_html(text: str) -> Tuple[str, List[str]]:
    """Strip HTML tags and extract text content.

    Uses Python's built-in html.parser — no external deps.

    Args:
        text: HTML string.

    Returns:
        Tuple of (extracted_text, error_messages).
    """
    errors: List[str] = []

    try:
        from html.parser import HTMLParser

        class _TextExtractor(HTMLParser):
            def __init__(self) -> None:
                super().__init__()
                self._parts: List[str] = []
                self._skip = False

            def handle_starttag(self, tag: str, attrs: List[Any]) -> None:
                if tag in ("script", "style", "noscript"):
                    self._skip = True

            def handle_endtag(self, tag: str) -> None:
                if tag in ("script", "style", "noscript"):
                    self._skip = False
                if tag in ("p", "div", "br", "h1", "h2", "h3", "h4", "li", "tr"):
                    self._parts.append("\n")

            def handle_data(self, data: str) -> None:
                if not self._skip and data.strip():
                    self._parts.append(data.strip())

        extractor = _TextExtractor()
        extractor.feed(text)
        result = " ".join(extractor._parts)
        result = result.replace(" \n ", "\n").replace("  ", " ").strip()
        return (result, errors)

    except Exception as e:
        errors.append(f"HTML parse error: {e}")
        return ("", errors)


# ──────────────────────────────────────────────────────────────
#  DocumentParser — Main public class
# ──────────────────────────────────────────────────────────────

class DocumentParser:
    """Python document parser for PDF, DOCX, and HTML files.

    Handles binary format parsing that the Rust ingestion engine
    cannot process natively. For TXT/CSV/JSON/Markdown, delegates
    to the Rust ``ingest_extract_text_simple`` function.

    Usage::

        parser = DocumentParser()

        # Parse a file by path
        text, errors = parser.parse_file("/path/to/document.pdf")

        # Parse bytes (e.g., from upload)
        text, errors = parser.parse_bytes("report.docx", docx_bytes)

        # Check if a format is supported
        if parser.supports("pdf"):
            ...
    """

    def parse_file(self, file_path: str) -> Tuple[str, List[str]]:
        """Parse a document file and extract its text content.

        Args:
            file_path: Absolute path to the document file.

        Returns:
            Tuple of (extracted_text, error_messages).
            Text may be empty if parsing fails; check errors list.
        """
        if not file_path or not isinstance(file_path, str):
            return ("", ["file_path must be a non-empty string"])

        if not os.path.exists(file_path):
            return ("", [f"File not found: {file_path}"])

        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return ("", ["File is empty"])
        if file_size > MAX_FILE_SIZE:
            return ("", [f"File too large: {file_size} bytes (max {MAX_FILE_SIZE})"])

        ext = self._get_extension(file_path)

        if ext == "pdf":
            return _parse_pdf(file_path)
        elif ext in ("docx", "doc"):
            return _parse_docx(file_path)
        elif ext in ("html", "htm"):
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    html_content = f.read()
                return _parse_html(html_content)
            except Exception as e:
                return ("", [f"HTML file read error: {e}"])
        elif ext in ("txt", "csv", "tsv", "json", "md", "markdown"):
            # Simple text formats — read and return
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    return (f.read(), [])
            except Exception as e:
                return ("", [f"Text file read error: {e}"])
        else:
            return ("", [f"Unsupported format: .{ext}"])

    def parse_bytes(
        self, filename: str, data: bytes
    ) -> Tuple[str, List[str]]:
        """Parse document bytes and extract text content.

        Args:
            filename: Original filename (used for format detection).
            data: Raw document bytes.

        Returns:
            Tuple of (extracted_text, error_messages).
        """
        if not filename or not isinstance(filename, str):
            return ("", ["filename must be a non-empty string"])
        if not data:
            return ("", ["Document data is empty"])
        if len(data) > MAX_FILE_SIZE:
            return ("", [f"Data too large: {len(data)} bytes (max {MAX_FILE_SIZE})"])

        ext = self._get_extension(filename)

        if ext == "pdf":
            return _parse_pdf_bytes(data)
        elif ext in ("docx", "doc"):
            return _parse_docx_bytes(data)
        elif ext in ("html", "htm"):
            try:
                html_content = data.decode("utf-8", errors="replace")
                return _parse_html(html_content)
            except Exception as e:
                return ("", [f"HTML decode error: {e}"])
        elif ext in ("txt", "csv", "tsv", "json", "md", "markdown"):
            try:
                text = data.decode("utf-8", errors="replace")
                return (text, [])
            except Exception as e:
                return ("", [f"Text decode error: {e}"])
        else:
            return ("", [f"Unsupported format: .{ext}"])

    def supports(self, format_str: str) -> bool:
        """Check if a format is supported for Python-side parsing.

        Args:
            format_str: Format string (e.g., ``"pdf"``, ``"docx"``).

        Returns:
            True if the format can be parsed by this module.
        """
        return format_str.lower() in PYTHON_FORMATS or format_str.lower() in (
            "txt", "csv", "tsv", "json", "md", "markdown", "html", "htm", "doc",
        )

    def available_parsers(self) -> Dict[str, bool]:
        """Check which Python parsers are actually available.

        Returns:
            Dict mapping parser name to availability status.
        """
        available: Dict[str, bool] = {}
        try:
            import PyPDF2  # noqa: F401
            available["pdf_pypdf2"] = True
        except ImportError:
            available["pdf_pypdf2"] = False

        try:
            import pdfminer  # noqa: F401
            available["pdf_pdfminer"] = True
        except ImportError:
            available["pdf_pdfminer"] = False

        try:
            import docx  # noqa: F401
            available["docx_python_docx"] = True
        except ImportError:
            available["docx_python_docx"] = False

        # HTML parser is always available (built-in)
        available["html_builtin"] = True

        return available

    @staticmethod
    def _get_extension(filename: str) -> str:
        """Extract and normalize file extension.

        Args:
            filename: Filename string.

        Returns:
            Lowercase extension without the dot, or empty string.
        """
        if not filename:
            return ""
        parts = filename.rsplit(".", 1)
        if len(parts) < 2:
            return ""
        return parts[1].lower()
