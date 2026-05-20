"""
Document Parser — Format-specific parsers.

PDF, DOCX, and HTML parsing functions for binary format support.
"""

from __future__ import annotations

import logging
import os
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────
#  PDF Parser
# ──────────────────────────────────────────────────────────────

def _parse_pdf(file_path: str) -> Tuple[str, List[str]]:
    """Extract text from a PDF file using Python libraries.

    Tries PyPDF2 first, then falls back to pdfminer.six.
    Returns (text, errors) tuple — never raises.
    """
    errors: List[str] = []
    text_parts: List[str] = []

    # Try PyPDF2
    try:
        from PyPDF2 import PdfReader  # type: ignore[import-untyped]

        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        if text_parts:
            return ("\n\n".join(text_parts), errors)
    except ImportError:
        errors.append("PyPDF2 not installed, trying pdfminer.six")
    except Exception as e:
        errors.append(f"PyPDF2 error: {e}")

    # Try pdfminer.six
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract  # type: ignore[import-untyped]

        text = pdfminer_extract(file_path)
        if text and text.strip():
            return (text, errors)
    except ImportError:
        errors.append("pdfminer.six not installed")
    except Exception as e:
        errors.append(f"pdfminer.six error: {e}")

    # No parser available
    if not text_parts:
        errors.append(
            "No PDF parser available. Install PyPDF2 or pdfminer.six: "
            "pip install PyPDF2 pdfminer.six"
        )

    return ("", errors)


def _parse_pdf_bytes(data: bytes) -> Tuple[str, List[str]]:
    """Extract text from PDF bytes.

    Writes bytes to a temporary file and delegates to _parse_pdf.
    """
    import tempfile

    errors: List[str] = []
    tmp_path: Optional[str] = None

    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        with open(tmp_path, "wb") as f:
            f.write(data)
        return _parse_pdf(tmp_path)
    except Exception as e:
        errors.append(f"PDF bytes parse error: {e}")
        return ("", errors)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


# ──────────────────────────────────────────────────────────────
#  DOCX Parser
# ──────────────────────────────────────────────────────────────

def _parse_docx(file_path: str) -> Tuple[str, List[str]]:
    """Extract text from a DOCX file using python-docx."""
    errors: List[str] = []
    text_parts: List[str] = []

    try:
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(file_path)

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_parts: List[str] = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_parts.append(cell.text.strip())
                if row_parts:
                    text_parts.append(" | ".join(row_parts))

        return ("\n".join(text_parts), errors)

    except ImportError:
        errors.append(
            "python-docx not installed. Install with: pip install python-docx"
        )
        return ("", errors)
    except Exception as e:
        errors.append(f"python-docx error: {e}")
        return ("", errors)


def _parse_docx_bytes(data: bytes) -> Tuple[str, List[str]]:
    """Extract text from DOCX bytes."""
    import tempfile

    errors: List[str] = []
    tmp_path: Optional[str] = None

    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        with open(tmp_path, "wb") as f:
            f.write(data)
        return _parse_docx(tmp_path)
    except Exception as e:
        errors.append(f"DOCX bytes parse error: {e}")
        return ("", errors)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


# ──────────────────────────────────────────────────────────────
#  HTML Parser
# ──────────────────────────────────────────────────────────────

def _parse_html(text: str) -> Tuple[str, List[str]]:
    """Strip HTML tags and extract text content.

    Uses Python's built-in html.parser — no external deps.
    """
    errors: List[str] = []

    try:
        from html.parser import HTMLParser

        class _TextExtractor(HTMLParser):
            def __init__(self) -> None:
                super().__init__()
                self._parts: List[str] = []
                self._skip = False

            def handle_starttag(self, tag: str, attrs: List[Any]) -> None:
                if tag in ("script", "style", "noscript"):
                    self._skip = True

            def handle_endtag(self, tag: str) -> None:
                if tag in ("script", "style", "noscript"):
                    self._skip = False
                if tag in ("p", "div", "br", "h1", "h2", "h3", "h4", "li", "tr"):
                    self._parts.append("\n")

            def handle_data(self, data: str) -> None:
                if not self._skip and data.strip():
                    self._parts.append(data.strip())

        extractor = _TextExtractor()
        extractor.feed(text)
        result = " ".join(extractor._parts)
        result = result.replace(" \n ", "\n").replace("  ", " ").strip()
        return (result, errors)

    except Exception as e:
        errors.append(f"HTML parse error: {e}")
        return ("", errors)
